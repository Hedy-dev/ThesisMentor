import uuid
import re
from dataclasses import dataclass, field, asdict
from typing import List, Optional, Dict
from docx import Document

class NodeType:
    ROOT = "ROOT"
    META = "SECTION_META"
    BODY = "BODY"
    SECTION = "SECTION"
    SUBSECTION = "SUBSECTION"
    PARAGRAPH = "PARAGRAPH"

@dataclass
class ThesisNode:
    id: str
    title: str
    type: str
    level: int
    parent_id: Optional[str] = None
    children: List[str] = field(default_factory=list)
    raw_text: Optional[str] = None
    clean_text: Optional[str] = None
    metadata: Dict = field(default_factory=dict)

def clean_text(text: str) -> str:
    text = re.sub(r'\d+\.', '', text)
    text = re.sub(r'\s+', ' ', text)
    text = text.lower().strip()
    return text

class ThesisParser:
    def __init__(self):
        self.nodes: Dict[str, ThesisNode] = {}
        self.root_id: Optional[str] = None
        self.current_stack: List[str] = []
        self.body_node_id: Optional[str] = None
        self.skip_content_section = False
        self.in_references_section = False

    def _generate_id(self):
        return str(uuid.uuid4())

    def _add_node(self, title, node_type, level, parent_id=None, raw_text=None, clean=None):
        node_id = self._generate_id()
        node = ThesisNode(
            id=node_id, title=title, type=node_type, level=level,
            parent_id=parent_id, raw_text=raw_text, clean_text=clean
        )
        self.nodes[node_id] = node
        if parent_id:
            self.nodes[parent_id].children.append(node_id)
        return node_id

    def _is_caps_title(self, text):
        return text.isupper() and len(text) > 3

    def _is_content_skip(self, text):
        return text.strip().upper() == "СОДЕРЖАНИЕ"

    def _is_references_section(self, text):
        keywords = ["СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", "СПИСОК ЛИТЕРАТУРЫ", "БИБЛИОГРАФИЯ"]
        return text.strip().upper() in keywords

    def parse(self, file_source) -> Dict:
        """
        Основной метод для интеграции. 
        file_source: путь к файлу или io.BytesIO
        """
        # Сбрасываем состояние перед парсингом нового файла
        self.nodes = {}
        self.current_stack = []
        self.body_node_id = None
        self.skip_content_section = False
        self.in_references_section = False

        doc = Document(file_source)
        
        # Поиск ROOT (Бизнес-тезис)
        found_root = False
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text.startswith("Бизнес-тезис"):
                self.root_id = self._add_node(title=text, node_type=NodeType.ROOT, level=0)
                self.current_stack = [self.root_id]
                found_root = True
                break
        
        if not found_root:
            self.root_id = self._add_node(title="Бизнес-тезис", node_type=NodeType.ROOT, level=0)
            self.nodes[self.root_id].metadata["warning"] = "Заголовок 'Бизнес-тезис' не найден"
            self.current_stack = [self.root_id]

        # Основной цикл
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text or text.startswith("Бизнес-тезис"):
                continue
            
            style = paragraph.style.name

            # Логика пропуска содержания
            if self._is_content_skip(text):
                self.skip_content_section = True
                continue
            if self.skip_content_section:
                if text.upper().startswith("ВВЕДЕНИЕ"):
                    self.skip_content_section = False
                else:
                    continue

            # Список источников
            if self._is_references_section(text):
                parent = self.body_node_id if self.body_node_id else self.root_id
                node_id = self._add_node(title=text, node_type=NodeType.SECTION, level=2, parent_id=parent)
                self.current_stack = [self.root_id, parent, node_id]
                self.in_references_section = True
                continue

            if self.in_references_section:
                continue

            # Создание BODY при встрече Введения
            if text.upper().startswith("ВВЕДЕНИЕ") and not self.body_node_id:
                self.body_node_id = self._add_node(title="Текст ВКР", node_type=NodeType.BODY, level=1, parent_id=self.root_id)
                self.current_stack = [self.root_id, self.body_node_id]

            # META (Heading 1)
            if style == "Heading 1":
                node_id = self._add_node(title=text, node_type=NodeType.META, level=1, parent_id=self.root_id)
                self.current_stack = [self.root_id, node_id]
                continue

            # SECTION (Heading 2 или CAPS)
            if style == "Heading 2" or self._is_caps_title(text):
                parent = self.body_node_id if self.body_node_id else self.root_id
                node_id = self._add_node(title=text, node_type=NodeType.SECTION, level=2, parent_id=parent)
                self.current_stack = [self.root_id, parent, node_id]
                continue

            # SUBSECTION (Heading 3)
            if style == "Heading 3":
                parent = self.current_stack[-1]
                node_id = self._add_node(title=text, node_type=NodeType.SUBSECTION, level=3, parent_id=parent)
                self.current_stack.append(node_id)
                continue

            # PARAGRAPH
            if style not in ["Heading 1", "Heading 2", "Heading 3"]:
                if not self.current_stack:
                    continue
                parent = self.current_stack[-1]
                self._add_node(
                    title="paragraph", node_type=NodeType.PARAGRAPH, level=4,
                    parent_id=parent, raw_text=text, clean=clean_text(text)
                )

        # Возвращаем структуру, конвертируя dataclasses в словари для совместимости с JSON/Критиком
        return {
            "root_id": self.root_id,
            "nodes": {node_id: asdict(node) for node_id, node in self.nodes.items()}
        }
